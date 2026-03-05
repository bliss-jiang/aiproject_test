from docx import Document
import os.path
from fastapi import UploadFile
import shutil
from fastapi import HTTPException
from typing import List, Tuple
from PyPDF2 import PdfReader


def extract_pdf_text_with_page_numbers(pdf_path) -> Tuple[str, List[int]]:
    text = ""
    page_numbers = []  # 用于存储包含文本的页码
    pdf_reader = PdfReader(pdf_path)  # 创建PDF阅读器对象
    for i, page in enumerate(pdf_reader.pages):  # 遍历每一页
        page_text = page.extract_text()  # 提取当前页的文本
        if page_text:  # 如果当前页有文本内容
            text += page_text  # 将文本添加到结果中
            page_numbers.append(i + 1)  # 记录页码(从1开始)
    print(f"extract_text_with_page_numbers函数 提取的文本长度: {len(text)} 个字符。")
    return text, page_numbers  # 返回提取的文本和页码列表

def get_file_txt_content(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        print(f"错误：文件 {file_path} 不存在")
        return None
    except UnicodeDecodeError:
        # 如果utf-8解码失败，尝试其他编码
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                content = file.read()
            return content
        except Exception as e:
            print(f"解码错误：{str(e)}")
            return None
    except Exception as e:
        print(f"读取文件时发生错误：{str(e)}")
        return None



def get_file_docx_content(file_path):
    try:
        doc = Document(file_path)
        full_text = []

        # 读取段落内容
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 读取表格内容（可选）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e:
        print(f"读取文件出错: {str(e)}")
        return None


def save_upload_file(upload_file: UploadFile) :
    """
    将上传的文件保存到指定路径
    """
    print("上传文件中")
    att_path = "attachments\\"+upload_file.filename
    print(att_path)
    if os.path.exists(att_path):
        return {"status": "existed", "file_name": att_path }
    try:
        with open(att_path,"wb+") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)

        upload_file.file.close()
    finally:
        return {"status": "successfully", "file_name": att_path }



def save_upload_file_doc(
        upload_file: UploadFile,
        max_size: int = 10 * 1024 * 1024,  # 默认最大10MB
) -> dict:
    # 检查文件大小
    file_size = 0
    for chunk in upload_file.file:
        file_size += len(chunk)
        if file_size > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"文件大小超过限制。最大允许大小: {max_size / (1024 * 1024)} MB"
            )

    # 重置文件指针
    upload_file.file.seek(0)

    # 处理文件名
    original_filename = upload_file.filename
    filename = original_filename

    # 创建目标文件路径
    file_path = "attachments\\" + filename

    # 保存文件
    try:
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
    except Exception as e:
        # 如果保存失败，删除可能已创建的文件
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")

    # 返回文件信息
    return {
        "filename": filename,
        "original_filename": original_filename,
        "file_path": str(file_path.absolute()),
        "file_size": file_size,
        "content_type": upload_file.content_type,
    }
